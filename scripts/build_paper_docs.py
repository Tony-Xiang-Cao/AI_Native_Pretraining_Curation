#!/usr/bin/env python3
"""Render paper/AutoCurate_AI_Native_Pretraining_Data.md to .docx and .pdf.

One Markdown parser feeds two renderers, so the Word and PDF outputs stay in
sync. Handles the subset of Markdown the paper uses: headings, paragraphs,
bold/italic/inline-code/links, bullet and (multi-line) numbered lists, fenced
code blocks, GitHub tables, block quotes, horizontal rules, and figure images.

Figures are SVG. On macOS we rasterize them with the built-in ``qlmanage``
(WebKit, full fidelity) and trim the surrounding whitespace before embedding;
if conversion is unavailable the figure is replaced by a caption pointer.

The PDF uses the Unicode DejaVu fonts bundled with matplotlib, so σ/β/Θ/Π/→/≈/−
and the other math glyphs render correctly (reportlab's built-in Type-1 fonts
can't).

Dependencies (the ``docx`` / ``pdf`` extras): python-docx, Pillow, reportlab,
matplotlib (for the bundled fonts only); macOS ``qlmanage`` for the figures.
"""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PAPER_MD = os.path.join(ROOT, "paper", "AutoCurate_AI_Native_Pretraining_Data.md")
PAPER_DOCX = os.path.join(ROOT, "paper", "AutoCurate_AI_Native_Pretraining_Data.docx")
PAPER_PDF = os.path.join(ROOT, "paper", "AutoCurate_AI_Native_Pretraining_Data.pdf")

TITLE = ("AutoCurate: An AI-Native, Self-Verifying Operating Loop "
         "for Pretraining-Data Curation")
AUTHOR = 'Xiang Tony Cao'

# A scratch dir for rasterized figures, populated by main() and read by both
# renderers. Keyed by the markdown's relative image path.
_FIG_CACHE: dict[str, str | None] = {}


# =============================================================================
# Figure conversion (SVG -> trimmed PNG) via macOS qlmanage
# =============================================================================
def svg_to_png(svg_path: str, out_dir: str) -> str | None:
    if not shutil.which("qlmanage"):
        return None
    subprocess.run(["qlmanage", "-t", "-s", "2000", "-o", out_dir, svg_path],
                   stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=False)
    raw = os.path.join(out_dir, os.path.basename(svg_path) + ".png")
    if not os.path.exists(raw):
        return None
    from PIL import Image, ImageChops
    img = Image.open(raw).convert("RGBA")
    # Composite onto white, then trim the uniform border qlmanage pads on.
    bg = Image.new("RGB", img.size, (255, 255, 255))
    bg.paste(img, mask=img.split()[3])
    diff = ImageChops.difference(bg, Image.new("RGB", bg.size, (255, 255, 255)))
    box = diff.getbbox()
    if box:
        pad = 12
        box = (max(0, box[0] - pad), max(0, box[1] - pad),
               min(bg.width, box[2] + pad), min(bg.height, box[3] + pad))
        bg = bg.crop(box)
    trimmed = os.path.join(out_dir, os.path.basename(svg_path) + ".trim.png")
    bg.save(trimmed)
    return trimmed


def resolve_image(rel: str) -> str | None:
    """Return a raster path for a markdown image ref (SVG -> PNG), or None."""
    if rel in _FIG_CACHE:
        return _FIG_CACHE[rel]
    src = os.path.normpath(os.path.join(os.path.dirname(PAPER_MD), rel))
    out: str | None = None
    if os.path.exists(src):
        if src.lower().endswith(".svg"):
            out = svg_to_png(src, _SCRATCH_DIR)
        else:
            out = src
    _FIG_CACHE[rel] = out
    return out


# =============================================================================
# Markdown -> blocks (a tiny, paper-specific parser shared by both renderers)
# =============================================================================
def parse_blocks(md: str):
    """Return a list of (kind, payload) blocks."""
    lines = md.splitlines()
    blocks = []
    i, n = 0, len(lines)
    seen_title = False
    while i < n:
        line = lines[i]

        if line.strip().startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", "\n".join(buf)))
            continue

        if line.lstrip().startswith("|") and i + 1 < n and re.match(
                r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            rows = []
            while i < n and lines[i].lstrip().startswith("|"):
                rows.append(lines[i])
                i += 1
            cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
            cells = [cells[0]] + cells[2:]  # drop |---| separator row
            blocks.append(("table", cells))
            continue

        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if m:
            blocks.append(("image", (m.group(1), m.group(2))))
            i += 1
            continue

        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            seen_title = True
            i += 1
            continue

        if line.strip() == "---":
            blocks.append(("hr", None))
            i += 1
            continue

        if line.lstrip().startswith(">"):
            buf = []
            while i < n and lines[i].lstrip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            blocks.append(("quote", " ".join(x for x in buf if x.strip())))
            continue

        # bullet list item (gather indented continuation lines)
        if re.match(r"^\s*[-*]\s+", line):
            buf = [re.sub(r"^\s*[-*]\s+", "", line)]
            i += 1
            while i < n and lines[i].strip() and lines[i][:1] in " \t" \
                    and not re.match(r"^\s*([-*]|\d+\.)\s+", lines[i]):
                buf.append(lines[i].strip())
                i += 1
            blocks.append(("bullet", " ".join(buf)))
            continue

        # numbered list item (keep the literal number; gather continuations)
        mnum = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if mnum:
            buf = [mnum.group(2)]
            i += 1
            while i < n and lines[i].strip() and lines[i][:1] in " \t" \
                    and not re.match(r"^\s*([-*]|\d+\.)\s+", lines[i]):
                buf.append(lines[i].strip())
                i += 1
            blocks.append(("number", (mnum.group(1), " ".join(buf))))
            continue

        if not line.strip():
            i += 1
            continue

        # plain paragraph: gather until blank / block start
        buf = [line]
        i += 1
        while i < n and lines[i].strip() and not re.match(
                r"^(\s*[-*]\s+|\s*\d+\.\s+|#{1,3}\s|>|\||```|!\[)", lines[i]) \
                and lines[i].strip() != "---":
            buf.append(lines[i])
            i += 1
        # The lines immediately under the title are the byline; keep their
        # breaks and centre them instead of collapsing into a body paragraph.
        if seen_title and not any(k not in ("h1",) for k, _ in blocks):
            blocks.append(("byline", [b.strip() for b in buf]))
        else:
            blocks.append(("para", " ".join(b.strip() for b in buf)))
    return blocks


_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*|\[[^\]]+\]\([^)]+\))")


def inline_tokens(text: str, styles: tuple = ()):
    """Split text into (styles, text) tokens, recursing into nested spans.

    ``styles`` is a tuple of the styles in force from the enclosing spans, so
    italic-inside-bold (``**... *x* ...**``) or code-inside-bold
    (``**...`x`...**``) compose correctly instead of leaking literal markers.
    """
    out = []
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) >= 4:
            out.extend(inline_tokens(tok[2:-2], styles + ("bold",)))
        elif tok.startswith("`") and tok.endswith("`"):
            out.append((styles + ("code",), tok[1:-1]))
        elif tok.startswith("*") and tok.endswith("*") and len(tok) >= 2:
            out.extend(inline_tokens(tok[1:-1], styles + ("italic",)))
        elif tok.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            out.append((styles + ("link",), m.group(1) if m else tok))
        else:
            out.append((styles, tok))
    return out


# =============================================================================
# DOCX renderer
# =============================================================================
def build_docx(blocks):
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    CODE_FONT, CODE_SHADE = "Consolas", "F2F2F2"

    def add_runs(paragraph, text):
        for styles, t in inline_tokens(text):
            r = paragraph.add_run(t)
            if "bold" in styles:
                r.bold = True
            if "italic" in styles:
                r.italic = True
            if "code" in styles:
                r.font.name = CODE_FONT
                r.font.size = Pt(9.5)

    def shade(el, fill):
        pr = el.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        pr.append(shd)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.9)
        section.left_margin = section.right_margin = Inches(1.0)
    for lvl in ("Heading 1", "Heading 2", "Heading 3", "Title"):
        try:
            doc.styles[lvl].font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        except KeyError:
            pass

    for kind, payload in blocks:
        if kind == "h1":
            doc.add_heading(payload, level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "byline":
            for ln in payload:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(p, ln)
        elif kind == "h2":
            doc.add_heading(payload, level=1)
        elif kind == "h3":
            doc.add_heading(payload, level=2)
        elif kind == "para":
            add_runs(doc.add_paragraph(), payload)
        elif kind == "bullet":
            add_runs(doc.add_paragraph(style="List Bullet"), payload)
        elif kind == "number":
            num, text = payload
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            p.add_run(f"{num}. ")
            add_runs(p, text)
        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            shade(p._p, "FBF3D6")
            add_runs(p, payload)
        elif kind == "code":
            p = doc.add_paragraph()
            shade(p._p, CODE_SHADE)
            run = p.add_run(payload)
            run.font.name = CODE_FONT
            run.font.size = Pt(9)
        elif kind == "hr":
            doc.add_paragraph()
        elif kind == "image":
            _, rel = payload
            png = resolve_image(rel)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if png:
                p.add_run().add_picture(png, width=Inches(6.0))
            else:
                p.add_run(f"[figure: {rel}]").italic = True
        elif kind == "table":
            cells = payload
            ncol = len(cells[0])
            table = doc.add_table(rows=0, cols=ncol)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r_idx, row in enumerate(cells):
                tr = table.add_row().cells
                for c_idx in range(ncol):
                    cp = tr[c_idx].paragraphs[0]
                    cp.paragraph_format.space_after = Pt(2)
                    cp.paragraph_format.space_before = Pt(2)
                    add_runs(cp, row[c_idx] if c_idx < len(row) else "")
                    if r_idx == 0:
                        for run in cp.runs:
                            run.bold = True
            doc.add_paragraph()

    doc.core_properties.title = TITLE
    doc.core_properties.author = AUTHOR
    doc.save(PAPER_DOCX)
    print(f"wrote {os.path.relpath(PAPER_DOCX, ROOT)}")


# =============================================================================
# PDF renderer (reportlab + DejaVu Unicode fonts)
# =============================================================================
def _register_fonts():
    import matplotlib
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    d = os.path.join(os.path.dirname(matplotlib.__file__), "mpl-data", "fonts", "ttf")
    fonts = {
        "DejaVu": "DejaVuSans.ttf",
        "DejaVu-Bold": "DejaVuSans-Bold.ttf",
        "DejaVu-Italic": "DejaVuSans-Oblique.ttf",
        "DejaVu-BoldItalic": "DejaVuSans-BoldOblique.ttf",
        "DejaVuMono": "DejaVuSansMono.ttf",
        "DejaVuMono-Bold": "DejaVuSansMono-Bold.ttf",
    }
    for name, fn in fonts.items():
        pdfmetrics.registerFont(TTFont(name, os.path.join(d, fn)))
    from reportlab.pdfbase.pdfmetrics import registerFontFamily
    registerFontFamily("DejaVu", normal="DejaVu", bold="DejaVu-Bold",
                       italic="DejaVu-Italic", boldItalic="DejaVu-BoldItalic")
    # A mono family too, so code-inside-bold (e.g. **`judgecurate`**) renders bold.
    registerFontFamily("DejaVuMono", normal="DejaVuMono", bold="DejaVuMono-Bold",
                       italic="DejaVuMono", boldItalic="DejaVuMono-Bold")


def _xml_escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline_to_rl(text: str) -> str:
    """Markdown inline -> reportlab mini-markup (nested styles compose)."""
    out = []
    for styles, t in inline_tokens(text):
        e = _xml_escape(t)
        if "code" in styles:
            e = f'<font face="DejaVuMono" size="8.5">{e}</font>'
        if "bold" in styles:
            e = f"<b>{e}</b>"
        if "italic" in styles:
            e = f"<i>{e}</i>"
        out.append(e)
    return "".join(out)


def build_pdf(blocks):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    _register_fonts()
    INK = colors.HexColor("#1f2933")
    ss = getSampleStyleSheet()

    def style(name, **kw):
        base = dict(fontName="DejaVu", textColor=INK, leading=14, fontSize=10.5)
        base.update(kw)
        return ParagraphStyle(name, parent=ss["Normal"], **base)

    body = style("body", spaceAfter=6)
    byline = style("byline", alignment=TA_CENTER, spaceAfter=2, fontSize=11)
    h1 = style("h1", fontName="DejaVu-Bold", fontSize=18, leading=22, spaceBefore=6,
               spaceAfter=10, alignment=TA_CENTER)
    h2 = style("h2", fontName="DejaVu-Bold", fontSize=14, leading=18, spaceBefore=12,
               spaceAfter=5)
    h3 = style("h3", fontName="DejaVu-Bold", fontSize=11.5, leading=15, spaceBefore=8,
               spaceAfter=3)
    bullet = style("bullet", leftIndent=16, bulletIndent=4, spaceAfter=3)
    numbered = style("numbered", leftIndent=20, bulletIndent=2, spaceAfter=3)
    quote = style("quote", leftIndent=14, rightIndent=10, backColor=colors.HexColor("#FBF3D6"),
                  borderColor=colors.HexColor("#E6D9A8"), borderWidth=0.5,
                  borderPadding=6, spaceBefore=4, spaceAfter=8, fontSize=9.5, leading=13)
    code = style("code", fontName="DejaVuMono", fontSize=8, leading=10.5,
                 backColor=colors.HexColor("#F2F2F2"), borderPadding=6,
                 spaceBefore=2, spaceAfter=8)

    story = []
    usable_w = letter[0] - 2 * inch

    for kind, payload in blocks:
        if kind == "h1":
            story.append(Paragraph(_inline_to_rl(payload), h1))
        elif kind == "byline":
            for ln in payload:
                story.append(Paragraph(_inline_to_rl(ln), byline))
            story.append(Spacer(1, 6))
        elif kind == "h2":
            story.append(Paragraph(_inline_to_rl(payload), h2))
        elif kind == "h3":
            story.append(Paragraph(_inline_to_rl(payload), h3))
        elif kind == "para":
            story.append(Paragraph(_inline_to_rl(payload), body))
        elif kind == "bullet":
            story.append(Paragraph(_inline_to_rl(payload), bullet, bulletText="•"))
        elif kind == "number":
            num, text = payload
            story.append(Paragraph(_inline_to_rl(text), numbered, bulletText=f"{num}."))
        elif kind == "quote":
            story.append(Paragraph(_inline_to_rl(payload), quote))
        elif kind == "code":
            safe = _xml_escape(payload).replace("\n", "<br/>").replace(" ", "&nbsp;")
            story.append(Paragraph(safe, code))
        elif kind == "hr":
            story.append(Spacer(1, 6))
        elif kind == "image":
            _, rel = payload
            png = resolve_image(rel)
            if png:
                from PIL import Image as PILImage
                iw, ih = PILImage.open(png).size
                w = min(usable_w, 6.2 * inch)
                story.append(Spacer(1, 4))
                story.append(Image(png, width=w, height=w * ih / iw))
                story.append(Spacer(1, 2))
        elif kind == "table":
            cells = payload
            data = [[Paragraph(_inline_to_rl(c), style("cell", fontSize=8.5, leading=11))
                     for c in row] for row in cells]
            ncol = len(cells[0])
            # first column a bit wider, rest equal
            first = usable_w * 0.32 if ncol > 2 else usable_w * 0.5
            rest = (usable_w - first) / max(1, ncol - 1)
            widths = [first] + [rest] * (ncol - 1)
            t = Table(data, colWidths=widths, repeatRows=1)
            t.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), "DejaVu"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "DejaVu-Bold"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#f4f6f8")]),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd2d9")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(Spacer(1, 2))
            story.append(t)
            story.append(Spacer(1, 8))

    doc = SimpleDocTemplate(
        PAPER_PDF, pagesize=letter,
        topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        leftMargin=inch, rightMargin=inch,
        title=TITLE, author=AUTHOR,
    )
    doc.build(story)
    print(f"wrote {os.path.relpath(PAPER_PDF, ROOT)}")


_SCRATCH_DIR = ""


def main() -> None:
    global _SCRATCH_DIR
    with open(PAPER_MD, encoding="utf-8") as fh:
        blocks = parse_blocks(fh.read())
    with tempfile.TemporaryDirectory() as tmp:
        _SCRATCH_DIR = tmp
        build_docx(blocks)
        build_pdf(blocks)


if __name__ == "__main__":
    main()
